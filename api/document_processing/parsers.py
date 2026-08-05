"""File-format parsing with a MinerU-first PDF path and local fallbacks."""

from __future__ import annotations

import json
import os
import shutil
import subprocess
import tempfile
from pathlib import Path


MINERU_TIMEOUT_SECONDS = int(os.getenv("MINERU_TIMEOUT_SECONDS", "300"))
MINERU_BACKEND = os.getenv("MINERU_BACKEND", "pipeline")
MINERU_MODEL_SOURCE = os.getenv("MINERU_MODEL_SOURCE", "modelscope")


def parse_document_pages(file_bytes: bytes, filename: str) -> tuple[list[tuple[int, str]], str]:
    """Parse a supported upload and return ``(pages, parser_name)``."""
    ext = Path(filename).suffix.lower()
    if ext == ".pdf":
        pages = parse_pdf_with_mineru(file_bytes, filename)
        if pages:
            return pages, "mineru"
        return parse_pdf_pages(file_bytes), "pymupdf/pdfplumber-ocr-fallback"
    if ext in {".docx", ".doc"}:
        return [(0, parse_docx(file_bytes))], "python-docx"
    if ext in {".txt", ".md"}:
        return [(0, file_bytes.decode("utf-8", errors="ignore"))], "plain-text"
    if ext == ".csv":
        return [(0, parse_csv(file_bytes))], "pandas/csv"
    raise ValueError(f"不支持的文件格式：{ext}")


def parse_pdf_with_mineru(file_bytes: bytes, filename: str) -> list[tuple[int, str]]:
    """Prefer MinerU for structured PDFs; return empty to request a fallback."""
    mineru_bin = shutil.which("mineru")
    if not mineru_bin:
        print("[Multimodal] MinerU CLI 未安装，使用本地 PDF 解析降级路径")
        return []

    safe_name = Path(filename).name or "document.pdf"
    if Path(safe_name).suffix.lower() != ".pdf":
        safe_name = "document.pdf"

    with tempfile.TemporaryDirectory(prefix="alphastock_mineru_") as temp_dir:
        temp_root = Path(temp_dir)
        source_path = temp_root / safe_name
        output_dir = temp_root / "output"
        source_path.write_bytes(file_bytes)
        command = [mineru_bin, "-p", str(source_path), "-o", str(output_dir), "-b", MINERU_BACKEND, "-m", "auto", "-l", "ch"]
        try:
            completed = subprocess.run(
                command,
                capture_output=True,
                timeout=MINERU_TIMEOUT_SECONDS,
                check=False,
                env={**os.environ, "MINERU_MODEL_SOURCE": MINERU_MODEL_SOURCE},
            )
        except subprocess.TimeoutExpired:
            print(f"[Multimodal] MinerU 解析超时（>{MINERU_TIMEOUT_SECONDS}s），使用降级路径")
            return []
        except OSError as exc:
            print(f"[Multimodal] 无法启动 MinerU：{exc}")
            return []

        if completed.returncode != 0:
            raw_error = completed.stderr or completed.stdout or b""
            stderr = raw_error.decode("utf-8", errors="replace")[-500:]
            print(f"[Multimodal] MinerU 解析失败，使用降级路径：{stderr}")
            return []

        content_lists = list(output_dir.rglob("*_content_list_v2.json"))
        if content_lists:
            try:
                pages = mineru_content_list_to_pages(content_lists[0])
                if pages:
                    return pages
            except (OSError, json.JSONDecodeError, TypeError, ValueError) as exc:
                print(f"[Multimodal] MinerU 结构化结果读取失败：{exc}")

        markdown_files = list(output_dir.rglob("*.md"))
        if markdown_files:
            markdown = markdown_files[0].read_text(encoding="utf-8", errors="ignore")
            if markdown.strip():
                return [(0, markdown)]
        print("[Multimodal] MinerU 未产生可消费的 Markdown/内容列表，使用降级路径")
        return []


def mineru_content_list_to_pages(content_list_path: Path) -> list[tuple[int, str]]:
    """Convert MinerU ``content_list_v2.json`` to page text with headings."""
    data = json.loads(content_list_path.read_text(encoding="utf-8"))
    if not isinstance(data, list):
        return []

    pages: list[tuple[int, str]] = []
    for page_index, items in enumerate(data, start=1):
        if not isinstance(items, list):
            continue
        blocks: list[str] = []
        for item in items:
            if not isinstance(item, dict):
                continue
            block_type = item.get("type", "")
            if block_type in {"page_header", "page_footer", "page_number", "page_aside_text"}:
                continue
            content = item.get("content", {})
            text = mineru_content_to_text(content).strip()
            if not text:
                continue
            if block_type == "title":
                level = content.get("level", 1) if isinstance(content, dict) else 1
                blocks.append(f"{'#' * max(1, min(int(level or 1), 6))} {text}")
            elif block_type == "table":
                blocks.append(f"表格：\n{text}")
            else:
                blocks.append(text)
        if blocks:
            pages.append((page_index, "\n\n".join(blocks)))
    return pages


def mineru_content_to_text(value) -> str:
    """Read meaningful text fields without embedding layout metadata in RAG."""
    if isinstance(value, str):
        return value
    if isinstance(value, list):
        return "".join(mineru_content_to_text(item) for item in value)
    if isinstance(value, dict):
        preferred = (
            "content", "title_content", "paragraph_content", "table_body",
            "table_caption", "table_footnote", "text_content", "list_items",
            "math_content", "code_content",
        )
        parts = [mineru_content_to_text(value[key]) for key in preferred if key in value]
        return "\n".join(part for part in parts if part)
    return ""


def parse_pdf_pages(file_bytes: bytes) -> list[tuple[int, str]]:
    """Use PyMuPDF text/OCR, then pdfplumber as a fallback, retaining pages."""
    try:
        import fitz
        doc = fitz.open(stream=file_bytes, filetype="pdf")
        pages: list[tuple[int, str]] = []
        for page_number, page in enumerate(doc, start=1):
            text = page.get_text()
            if len(text.strip()) > 50:
                pages.append((page_number, text))
                continue
            try:
                import io
                import pytesseract
                from PIL import Image
                pix = page.get_pixmap(matrix=fitz.Matrix(2, 2))
                ocr_text = pytesseract.image_to_string(Image.open(io.BytesIO(pix.tobytes("png"))), lang="chi_sim+eng")
                if ocr_text.strip():
                    pages.append((page_number, ocr_text))
            except ImportError:
                if text.strip():
                    pages.append((page_number, text))
        return pages
    except ImportError:
        try:
            import io
            import pdfplumber
            with pdfplumber.open(io.BytesIO(file_bytes)) as pdf:
                pages: list[tuple[int, str]] = []
                for page_number, page in enumerate(pdf.pages, start=1):
                    page_parts = []
                    text = page.extract_text() or ""
                    for table in page.extract_tables():
                        for row in table:
                            if row:
                                page_parts.append(" | ".join(str(cell) for cell in row if cell))
                    if text:
                        page_parts.append(text)
                    if page_parts:
                        pages.append((page_number, "\n".join(page_parts)))
                return pages
        except Exception as exc:
            raise Exception(f"PDF解析失败（需要安装 PyMuPDF 或 pdfplumber）: {exc}") from exc


def parse_docx(file_bytes: bytes) -> str:
    try:
        import io
        from docx import Document
        doc = Document(io.BytesIO(file_bytes))
        paragraphs = [paragraph.text for paragraph in doc.paragraphs if paragraph.text.strip()]
        table_texts = [
            " | ".join(cell.text.strip() for cell in row.cells if cell.text.strip())
            for table in doc.tables for row in table.rows
        ]
        return "\n".join(paragraphs + [text for text in table_texts if text])
    except ImportError as exc:
        raise Exception("解析Word文档需要安装 python-docx：pip install python-docx") from exc


def parse_csv(file_bytes: bytes) -> str:
    try:
        import io
        import pandas as pd
        return pd.read_csv(io.BytesIO(file_bytes), encoding="utf-8-sig").to_string(index=False, max_rows=200)
    except Exception:
        return file_bytes.decode("utf-8-sig", errors="ignore")

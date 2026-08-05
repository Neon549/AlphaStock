#!/usr/bin/env python3
# -*- coding: utf-8 -*-

"""
@author: yulin
@created: 2026/7/17 21:20
@updated: 2026/7/17 21:20
@version: 1.0
@description: 
"""
#!/usr/bin/env python3
# -*- coding: utf-8 -*-

"""
api/multimodal.py
多模态处理器

支持：
  1. 图片分析（财报截图/K线图）→ Qwen-VL 提取数据
  2. PDF/文档解析 → 分块入 PostgreSQL + pgvector → RAG检索
  3. 用户级隔离：每个session独立Collection，对话结束清理

大厂方案：
  图片：多模态LLM直接理解，不OCR
  文档：解析→分块→临时向量库→检索增强
"""

import hashlib
import json
import os
import shutil
import subprocess
import tempfile
from datetime import datetime
from pathlib import Path
from typing import Optional
import re

from db import get_conn
from rag.news_indexer import _embed
from api.multimodal_vision import analyze_image
from api.document_processing.chunking import (
    CHILD_CHUNK_OVERLAP,
    CHILD_CHUNK_SIZE,
    build_hierarchical_chunks,
    parse_heading,
)
from api.document_processing.parsers import parse_document_pages


# ── 临时文档向量库（PostgreSQL + pgvector）────────────────────────────────

COLLECTION_TTL_HOURS = 2  # 2小时后自动清理

# ``process_document`` 已使用 document_processing.parsers；以下私有入口暂留，
# 只为可能存在的旧内部脚本提供过渡兼容，下一轮会随调用方审计一起移除。
MINERU_TIMEOUT_SECONDS = int(os.getenv("MINERU_TIMEOUT_SECONDS", "300"))
MINERU_BACKEND = os.getenv("MINERU_BACKEND", "pipeline")
MINERU_MODEL_SOURCE = os.getenv("MINERU_MODEL_SOURCE", "modelscope")

_EVIDENCE_ID_PATTERN = re.compile(r"evidence_id=([^|\]]+)")
_EVIDENCE_FIELD_PATTERN = re.compile(r"(文件|章节|版本)=([^|\]]+)")
_EVIDENCE_PAGE_PATTERN = re.compile(r"第\s*(\d+)\s*页")


def cleanup_session(session_id: str):
    """清理指定 session 的临时文档，不保留本地向量库副本。"""
    try:
        with get_conn() as conn:
            with conn.cursor() as cur:
                cur.execute("DELETE FROM uploaded_document_chunks WHERE session_id = %s", (session_id,))
                deleted = cur.rowcount
            conn.commit()
        print(f"[Multimodal] 清理 session {session_id[:8]} 的临时文档：{deleted} 块")
    except Exception as exc:
        print(f"[Multimodal] 清理临时文档失败: {exc}")


def cleanup_document(session_id: str, document_id: str) -> int:
    """只清理会话中的一份文档，供前端移除附件时调用。"""
    try:
        with get_conn() as conn:
            with conn.cursor() as cur:
                cur.execute(
                    "DELETE FROM uploaded_document_chunks WHERE session_id = %s AND document_id = %s",
                    (session_id, document_id),
                )
                deleted = cur.rowcount
            conn.commit()
        return deleted
    except Exception as exc:
        print(f"[Multimodal] 清理单个临时文档失败: {exc}")
        return 0


def cleanup_expired_sessions():
    """清理超过 TTL 的临时文档，返回删除块数。"""
    try:
        with get_conn() as conn:
            with conn.cursor() as cur:
                cur.execute(
                    "DELETE FROM uploaded_document_chunks "
                    "WHERE created_at < NOW() - (%s * INTERVAL '1 hour')",
                    (COLLECTION_TTL_HOURS,),
                )
                deleted = cur.rowcount
            conn.commit()
        return deleted
    except Exception as exc:
        print(f"[Multimodal] 清理过期临时文档失败: {exc}")
        return 0


def index_image_analysis(
    image_bytes: bytes,
    filename: str,
    session_id: str,
    vlm_result: dict,
    question: str = "",
) -> dict:
    """Store a successful VLM extraction as session-scoped pgvector evidence.

    The image itself is not embedded.  We embed the VLM-extracted text, while retaining
    the image filename and a dedicated parent path so later citations remain honest about
    their source.  Images do not have a reliable document page number and use page=0.
    """
    extracted_data = (vlm_result.get("extracted_data") or "").strip()
    data_type = vlm_result.get("data_type", "unknown")
    if not extracted_data or data_type == "unknown":
        return {"success": False, "indexed": False, "error": "VLM 未返回可索引的识别结果"}

    image_type_label = {
        "financial": "财报或财务截图",
        "kline": "K线或技术图表",
        "other": "公告、新闻或其他图片",
    }.get(data_type, "图片")
    source_text = (
        "# 图像识别结果\n"
        f"## 图像类型\n{image_type_label}\n"
        f"## 用户问题\n{question or '未提供'}\n"
        f"## VLM 提取事实\n{extracted_data}"
    )
    chunks = _build_hierarchical_chunks([(0, source_text)])
    if not chunks:
        return {"success": False, "indexed": False, "error": "图像识别结果分块失败"}

    document_id = hashlib.sha256(image_bytes).hexdigest()[:16]
    storage_id = hashlib.sha256(f"{session_id}:{document_id}".encode("utf-8")).hexdigest()[:16]
    ids = [f"img_{storage_id}_{index}" for index in range(len(chunks))]
    document_version = "qwen-vl-plus-extraction-v1"
    uploaded_at = datetime.now().isoformat()

    try:
        embeddings = _embed([chunk["text"] for chunk in chunks])
        with get_conn() as conn:
            with conn.cursor() as cur:
                # Re-uploading the same image replaces its prior extraction for this session.
                cur.execute(
                    "DELETE FROM uploaded_document_chunks WHERE session_id = %s AND document_id = %s",
                    (session_id, document_id),
                )
                for index, (chunk, embedding) in enumerate(zip(chunks, embeddings)):
                    cur.execute(
                        """
                        INSERT INTO uploaded_document_chunks
                            (id, session_id, filename, document_id, document_version,
                             chunk_index, page, parent_path, previous_id, next_id,
                             content, embedding, created_at)
                        VALUES
                            (%s, %s, %s, %s, %s, %s, %s, %s, %s, %s, %s, %s::vector, %s)
                        """,
                        (
                            ids[index], session_id, filename, document_id, document_version,
                            index, 0, chunk["parent_path"],
                            ids[index - 1] if index > 0 else None,
                            ids[index + 1] if index + 1 < len(ids) else None,
                            chunk["text"], str(embedding), uploaded_at,
                        ),
                    )
            conn.commit()
    except Exception as exc:
        return {"success": False, "indexed": False, "error": f"VLM 结果写入 pgvector 失败：{exc}"}

    return {
        "success": True,
        "indexed": True,
        "document_id": document_id,
        "chunk_count": len(chunks),
        "document_version": document_version,
    }


# ── 文档处理（PDF/Word/TXT）──────────────────────────────────────────────

def process_document(
    file_bytes: bytes,
    filename: str,
    session_id: str,
) -> dict:
    """
    处理上传的文档，解析后存入临时向量库

    Args:
        file_bytes: 文件二进制数据
        filename: 文件名（用于判断格式）
        session_id: 用户session ID（隔离用）

    Returns:
        {
            "success": True,
            "chunk_count": 15,
            "preview": "文档内容预览...",
            "file_type": "pdf"
        }
    """
    ext = Path(filename).suffix.lower()
    # 解析文档内容。PDF 保留页码；其它格式没有可靠页码时以 0 表示。
    try:
        pages, parser = parse_document_pages(file_bytes, filename)
    except Exception as e:
        return {"success": False, "error": f"文档解析失败：{e}"}

    text = "\n\n".join(page_text for _, page_text in pages if page_text.strip())
    if not text.strip():
        return {"success": False, "error": "文档内容为空"}

    document_id = hashlib.sha256(file_bytes).hexdigest()[:16]
    chunks = _build_hierarchical_chunks(pages)

    if not chunks:
        return {"success": False, "error": "文档分块失败"}

    # 子块用于向量召回；parent_path、页码和相邻子块 ID 用于命中后的上下文补全。
    # session 级文档量较小，因此不建 HNSW；查询先按 session 过滤后精确向量排序。
    uploaded_at = datetime.now().isoformat()
    # 主键必须同时隔离 session；同一文件可被不同用户/会话上传。
    storage_id = hashlib.sha256(f"{session_id}:{document_id}".encode("utf-8")).hexdigest()[:16]
    ids = [f"doc_{storage_id}_{i}" for i in range(len(chunks))]
    embeddings = _embed([chunk["text"] for chunk in chunks])
    try:
        with get_conn() as conn:
            with conn.cursor() as cur:
                # 重传同一份文件时替换它；同一 session 的其他附件仍应可检索。
                cur.execute(
                    "DELETE FROM uploaded_document_chunks WHERE session_id = %s AND document_id = %s",
                    (session_id, document_id),
                )
                for i, (chunk, embedding) in enumerate(zip(chunks, embeddings)):
                    cur.execute(
                        """
                        INSERT INTO uploaded_document_chunks
                            (id, session_id, filename, document_id, document_version,
                             chunk_index, page, parent_path, previous_id, next_id,
                             content, embedding, created_at)
                        VALUES
                            (%s, %s, %s, %s, %s, %s, %s, %s, %s, %s, %s, %s::vector, %s)
                        """,
                        (
                            ids[i], session_id, filename, document_id, document_id,
                            i, chunk["page"], chunk["parent_path"],
                            ids[i - 1] if i > 0 else None,
                            ids[i + 1] if i + 1 < len(ids) else None,
                            chunk["text"], str(embedding), uploaded_at,
                        ),
                    )
            conn.commit()
    except Exception as exc:
        return {"success": False, "error": f"临时文档写入 PostgreSQL 失败：{exc}"}

    return {
        "success": True,
        "chunk_count": len(chunks),
        "preview": text[:300] + "..." if len(text) > 300 else text,
        "file_type": ext.lstrip("."),
        "total_chars": len(text),
        "document_id": document_id,
        "retrieval_mode": "hierarchical_child_chunk_with_neighbor_context",
        "parser": parser,
    }


def retrieve_from_document(session_id: str, query: str, k: int = 5) -> str:
    """
    从用户上传的文档中检索相关内容

    Args:
        session_id: 用户session ID
        query: 检索问题
        k: 返回条数

    Returns:
        检索到的文档内容字符串
    """
    try:
        query_embedding = _embed([query])[0]
        with get_conn() as conn:
            with conn.cursor() as cur:
                cur.execute(
                    """
                    SELECT id, content, filename, document_version, page, parent_path,
                           previous_id, next_id
                    FROM uploaded_document_chunks
                    WHERE session_id = %s
                    ORDER BY embedding <=> %s::vector
                    LIMIT %s
                    """,
                    (session_id, str(query_embedding), k),
                )
                rows = cur.fetchall()
        if not rows:
            return ""
        rendered: list[str] = []
        included_ids: set[str] = set()
        for row in rows:
            chunk_id, chunk, *metadata_values = row
            metadata = _document_row_metadata(metadata_values)
            _append_evidence(rendered, included_ids, chunk_id, chunk, metadata, "命中子块")

            # 命中精确子块后，最多补一个相邻块。它能补齐被切分的表格说明或限定条件，
            # 又不会像直接塞整章那样明显拉高 token 成本。
            for neighbor_id in (metadata.get("previous_id"), metadata.get("next_id")):
                if not neighbor_id or neighbor_id in included_ids:
                    continue
                with get_conn() as conn:
                    with conn.cursor() as cur:
                        cur.execute(
                            """
                            SELECT id, content, filename, document_version, page, parent_path,
                                   previous_id, next_id
                            FROM uploaded_document_chunks
                            WHERE session_id = %s AND id = %s
                            """,
                            (session_id, neighbor_id),
                        )
                        neighbor = cur.fetchone()
                if neighbor:
                    neighbor_id, neighbor_text, *neighbor_metadata_values = neighbor
                    _append_evidence(
                        rendered,
                        included_ids,
                        neighbor_id,
                        neighbor_text,
                        _document_row_metadata(neighbor_metadata_values),
                        "相邻上下文",
                    )

        return "\n\n".join(rendered)

    except Exception as e:
        print(f"[Multimodal] 文档检索失败: {e}")
        return ""


def extract_document_citations(document_evidence: str) -> list[dict]:
    """从供模型使用的证据块提取可直接返回给前端的引用元数据。"""
    citations: list[dict] = []
    seen: set[str] = set()
    for header in re.findall(r"\[[^\]]*evidence_id=[^\]]+\]", document_evidence or ""):
        id_match = _EVIDENCE_ID_PATTERN.search(header)
        if not id_match:
            continue
        evidence_id = id_match.group(1).strip()
        if evidence_id in seen:
            continue
        seen.add(evidence_id)
        fields = {key: value.strip() for key, value in _EVIDENCE_FIELD_PATTERN.findall(header)}
        page_match = _EVIDENCE_PAGE_PATTERN.search(header)
        citations.append(
            {
                "evidence_id": evidence_id,
                "filename": fields.get("文件", ""),
                "section": fields.get("章节", "正文"),
                "page": int(page_match.group(1)) if page_match else None,
                "document_version": fields.get("版本", ""),
            }
        )
    return citations


# ── 文档解析器 ────────────────────────────────────────────────────────────

def _document_row_metadata(values: list) -> dict:
    """将 PostgreSQL 查询行转换为统一证据 metadata。"""
    filename, document_version, page, parent_path, previous_id, next_id = values
    return {
        "filename": filename,
        "document_version": document_version,
        "page": page,
        "parent_path": parent_path,
        "previous_id": previous_id,
        "next_id": next_id,
    }


def _append_evidence(
    rendered: list[str],
    included_ids: set[str],
    chunk_id: str,
    text: str,
    metadata: dict,
    label: str,
) -> None:
    """格式化可回链的检索证据，供下游 prompt 与最终报告引用。"""
    if chunk_id in included_ids:
        return
    included_ids.add(chunk_id)
    page = metadata.get("page", 0)
    page_text = f"第 {page} 页" if page else "无页码"
    rendered.append(
        f"[{label} | evidence_id={chunk_id} | 文件={metadata.get('filename', '')} | "
        f"章节={metadata.get('parent_path', '正文')} | {page_text} | "
        f"版本={metadata.get('document_version', '')}]\n{text}"
    )


_build_hierarchical_chunks = build_hierarchical_chunks
_parse_heading = parse_heading


def _parse_pdf_with_mineru(file_bytes: bytes, filename: str) -> list[tuple[int, str]]:
    """优先用 MinerU 把复杂 PDF 解析为保留结构和页码的内容块。

    MinerU 的 ``content_list_v2.json`` 以页面组织 title/paragraph/table 等块，
    因此比仅提取 PDF 文字层更适合财报、双栏版面和扫描件。解析失败时返回空列表，
    调用方会显式降级到现有 PyMuPDF/pdfplumber/OCR 路径。
    """
    mineru_bin = shutil.which("mineru")
    if not mineru_bin:
        print("[Multimodal] MinerU CLI 未安装，使用本地 PDF 解析降级路径")
        return []

    safe_name = Path(filename).name or "document.pdf"
    if Path(safe_name).suffix.lower() != ".pdf":
        safe_name = "document.pdf"

    with tempfile.TemporaryDirectory(prefix="alphastock_mineru_") as temp_dir:
        temp_root = Path(temp_dir)
        source_path = temp_root / safe_name
        output_dir = temp_root / "output"
        source_path.write_bytes(file_bytes)

        command = [
            mineru_bin,
            "-p", str(source_path),
            "-o", str(output_dir),
            "-b", MINERU_BACKEND,
            "-m", "auto",
            "-l", "ch",
        ]
        try:
            completed = subprocess.run(
                command,
                capture_output=True,
                timeout=MINERU_TIMEOUT_SECONDS,
                check=False,
                env={**os.environ, "MINERU_MODEL_SOURCE": MINERU_MODEL_SOURCE},
            )
        except subprocess.TimeoutExpired:
            print(f"[Multimodal] MinerU 解析超时（>{MINERU_TIMEOUT_SECONDS}s），使用降级路径")
            return []
        except OSError as exc:
            print(f"[Multimodal] 无法启动 MinerU：{exc}")
            return []

        if completed.returncode != 0:
            raw_error = completed.stderr or completed.stdout or b""
            stderr = raw_error.decode("utf-8", errors="replace")[-500:]
            print(f"[Multimodal] MinerU 解析失败，使用降级路径：{stderr}")
            return []

        content_lists = list(output_dir.rglob("*_content_list_v2.json"))
        if content_lists:
            try:
                pages = _mineru_content_list_to_pages(content_lists[0])
                if pages:
                    return pages
            except (OSError, json.JSONDecodeError, TypeError, ValueError) as exc:
                print(f"[Multimodal] MinerU 结构化结果读取失败：{exc}")

        # 极少数后端只产生 Markdown；仍能保留标题层级，但无法可靠标注页码。
        markdown_files = list(output_dir.rglob("*.md"))
        if markdown_files:
            markdown = markdown_files[0].read_text(encoding="utf-8", errors="ignore")
            if markdown.strip():
                return [(0, markdown)]

        print("[Multimodal] MinerU 未产生可消费的 Markdown/内容列表，使用降级路径")
        return []


def _mineru_content_list_to_pages(content_list_path: Path) -> list[tuple[int, str]]:
    """将 MinerU v3 ``content_list_v2.json`` 转成带 Markdown 标题的页级文本。"""
    data = json.loads(content_list_path.read_text(encoding="utf-8"))
    if not isinstance(data, list):
        return []

    pages: list[tuple[int, str]] = []
    for page_index, items in enumerate(data, start=1):
        if not isinstance(items, list):
            continue
        blocks: list[str] = []
        for item in items:
            if not isinstance(item, dict):
                continue
            block_type = item.get("type", "")
            # 页眉、页脚、页码属于版面噪声，不能进入 RAG 证据。
            if block_type in {"page_header", "page_footer", "page_number", "page_aside_text"}:
                continue
            content = item.get("content", {})
            text = _mineru_content_to_text(content).strip()
            if not text:
                continue
            if block_type == "title":
                level = content.get("level", 1) if isinstance(content, dict) else 1
                level = max(1, min(int(level or 1), 6))
                blocks.append(f"{'#' * level} {text}")
            elif block_type == "table":
                blocks.append(f"表格：\n{text}")
            else:
                blocks.append(text)
        if blocks:
            pages.append((page_index, "\n\n".join(blocks)))
    return pages


def _mineru_content_to_text(value) -> str:
    """递归兼容 MinerU 内容块中的 span/list/HTML 字段。"""
    if isinstance(value, str):
        return value
    if isinstance(value, list):
        return "".join(_mineru_content_to_text(item) for item in value)
    if isinstance(value, dict):
        # 先取最常见的内容字段，避免将 bbox、路径等元数据混入检索文本。
        preferred = (
            "content", "title_content", "paragraph_content", "table_body",
            "table_caption", "table_footnote", "text_content", "list_items",
            "math_content", "code_content",
        )
        parts = [_mineru_content_to_text(value[key]) for key in preferred if key in value]
        return "\n".join(part for part in parts if part)
    return ""


def _parse_pdf_pages(file_bytes: bytes) -> list[tuple[int, str]]:
    """解析 PDF，优先提取文字层，失败时 OCR；保留原始页码。"""
    try:
        import fitz  # PyMuPDF

        doc = fitz.open(stream=file_bytes, filetype="pdf")
        pages: list[tuple[int, str]] = []

        for page_number, page in enumerate(doc, start=1):
            text = page.get_text()
            if len(text.strip()) > 50:
                # 有文字层，直接用
                pages.append((page_number, text))
            else:
                # 扫描件，尝试OCR
                try:
                    import pytesseract
                    from PIL import Image
                    import io

                    mat = fitz.Matrix(2, 2)  # 放大2倍提高OCR精度
                    pix = page.get_pixmap(matrix=mat)
                    img = Image.open(io.BytesIO(pix.tobytes("png")))
                    ocr_text = pytesseract.image_to_string(img, lang="chi_sim+eng")
                    if ocr_text.strip():
                        pages.append((page_number, ocr_text))
                except ImportError:
                    # 没装pytesseract，跳过OCR
                    if text.strip():
                        pages.append((page_number, text))

        return pages

    except ImportError:
        # 没装PyMuPDF，尝试pdfplumber
        try:
            import pdfplumber
            import io

            with pdfplumber.open(io.BytesIO(file_bytes)) as pdf:
                pages: list[tuple[int, str]] = []
                for page_number, page in enumerate(pdf.pages, start=1):
                    page_parts = []
                    text = page.extract_text() or ""
                    # 提取表格
                    tables = page.extract_tables()
                    for table in tables:
                        for row in table:
                            if row:
                                page_parts.append(" | ".join(str(c) for c in row if c))
                    if text:
                        page_parts.append(text)
                    if page_parts:
                        pages.append((page_number, "\n".join(page_parts)))
                return pages
        except Exception as e:
            raise Exception(f"PDF解析失败（需要安装 PyMuPDF 或 pdfplumber）: {e}")


def _parse_docx(file_bytes: bytes) -> str:
    """解析Word文档"""
    try:
        from docx import Document
        import io

        doc = Document(io.BytesIO(file_bytes))
        paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]

        # 提取表格
        table_texts = []
        for table in doc.tables:
            for row in table.rows:
                row_text = " | ".join(cell.text.strip() for cell in row.cells if cell.text.strip())
                if row_text:
                    table_texts.append(row_text)

        return "\n".join(paragraphs + table_texts)

    except ImportError:
        raise Exception("解析Word文档需要安装 python-docx：pip install python-docx")


def _parse_csv(file_bytes: bytes) -> str:
    """解析CSV文件"""
    try:
        import pandas as pd
        import io

        df = pd.read_csv(io.BytesIO(file_bytes), encoding="utf-8-sig")
        return df.to_string(index=False, max_rows=200)
    except Exception as e:
        # 纯文本方式
        return file_bytes.decode("utf-8-sig", errors="ignore")
